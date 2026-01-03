# backend/notifications/document_generator.py
"""
Generador de documentos PDF y DOCX para rutinas de entrenamiento.
"""
from __future__ import annotations

from io import BytesIO
from typing import Dict, Any, List, Optional
from datetime import datetime

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_routine_pdf(routine_data: Dict[str, Any]) -> BytesIO:
    """
    Genera un PDF de la rutina de entrenamiento.
    
    Args:
        routine_data: Diccionario con la estructura de routine_detail del chatbot
        
    Returns:
        BytesIO con el contenido del PDF
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch
    )
    
    story = []
    styles = getSampleStyleSheet()
    
    # Estilo personalizado para el título
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1e40af'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Estilo para subtítulos
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#374151'),
        spaceAfter=10,
        fontName='Helvetica-Bold'
    )
    
    # Estilo para texto normal
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=6
    )
    
    # Título principal
    header = routine_data.get("header", "Rutina de Entrenamiento")
    story.append(Paragraph(header, title_style))
    story.append(Spacer(1, 0.2 * inch))
    
    # Información del resumen
    summary = routine_data.get("summary", {})
    summary_data = [
        ["Tiempo estimado:", f"{summary.get('tiempo_min', '-')} minutos"],
        ["Ejercicios:", str(summary.get('ejercicios', '-'))],
        ["Equipamiento:", summary.get('equipamiento', '-')],
        ["Objetivo:", summary.get('objetivo', '-')],
        ["Nivel:", summary.get('nivel', '-')],
        ["Músculo:", summary.get('musculo', '-')],
    ]
    
    summary_table = Table(summary_data, colWidths=[2 * inch, 4 * inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#111827')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#d1d5db')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 0.3 * inch))
    
    # Precauciones de salud
    health_notes = summary.get('health_notes')
    if health_notes and isinstance(health_notes, list) and len(health_notes) > 0:
        story.append(Paragraph("⚠️ Precauciones:", subtitle_style))
        for note in health_notes:
            story.append(Paragraph(f"• {note}", normal_style))
        story.append(Spacer(1, 0.2 * inch))
    
    # Alergias
    allergies = summary.get('allergies')
    if allergies and isinstance(allergies, list) and len(allergies) > 0:
        allergy_text = f"Alergias registradas: {', '.join(allergies)}"
        story.append(Paragraph(allergy_text, normal_style))
        story.append(Spacer(1, 0.2 * inch))
    
    # Fallback notice
    fallback_notice = routine_data.get('fallback_notice')
    if fallback_notice:
        story.append(Paragraph(f"ℹ️ {fallback_notice}", normal_style))
        story.append(Spacer(1, 0.2 * inch))
    
    # Ejercicios
    exercises = routine_data.get('exercises', [])
    if exercises:
        story.append(Paragraph("Detalle de Ejercicios", subtitle_style))
        story.append(Spacer(1, 0.1 * inch))
        
        for exercise in exercises:
            ex_name = exercise.get('nombre', 'Ejercicio')
            series = exercise.get('series', '-')
            reps = exercise.get('repeticiones', '-')
            rpe = exercise.get('rpe', '-')
            rir = exercise.get('rir', '-')
            video = exercise.get('video', '')
            
            # Nombre del ejercicio
            ex_title = ParagraphStyle(
                'ExerciseTitle',
                parent=styles['Normal'],
                fontSize=11,
                fontName='Helvetica-Bold',
                textColor=colors.HexColor('#1f2937'),
                spaceAfter=4
            )
            story.append(Paragraph(f"{exercise.get('orden', '')}. {ex_name}", ex_title))
            
            # Detalles del ejercicio
            details = f"Series: {series} • Repeticiones: {reps} • {rpe} • {rir}"
            story.append(Paragraph(details, normal_style))
            
            if video:
                video_text = f'<font color="#2563eb">🎥 Video: {video}</font>'
                story.append(Paragraph(video_text, normal_style))
            
            story.append(Spacer(1, 0.15 * inch))
    
    # Progresión
    progresion = summary.get('progresion')
    if progresion:
        story.append(Spacer(1, 0.1 * inch))
        prog_style = ParagraphStyle(
            'Progression',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#6b7280'),
            italic=True
        )
        story.append(Paragraph(f"📈 Progresión: {progresion}", prog_style))
    
    # Footer
    story.append(Spacer(1, 0.3 * inch))
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=colors.HexColor('#9ca3af'),
        alignment=TA_CENTER
    )
    footer_text = f"Generado por Fitter • {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    story.append(Paragraph(footer_text, footer_style))
    
    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_hero_plan_pdf(plan_data: Dict[str, Any]) -> BytesIO:
    """Genera un PDF para un entreno unico (hero plan)."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    story = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "HeroTitle",
        parent=styles["Heading1"],
        fontSize=18,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
    )
    subtitle_style = ParagraphStyle(
        "HeroSubtitle",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#334155"),
        spaceAfter=8,
        fontName="Helvetica-Bold",
    )
    normal_style = ParagraphStyle(
        "HeroNormal",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
    )

    plan_title = plan_data.get("title") or "Entreno unico"
    payload = plan_data.get("payload") or {}

    story.append(Paragraph(plan_title, title_style))
    story.append(Spacer(1, 0.2 * inch))

    meta_data = [
        ["Duracion:", payload.get("duration", "-")],
        ["Enfoque:", payload.get("focus", "-")],
        ["Tipo de cuerpo:", payload.get("body_type", "-")],
        ["Inicio:", payload.get("start", "-")],
        ["Equipamiento:", payload.get("equipment", "-")],
    ]
    meta_table = Table(meta_data, colWidths=[1.6 * inch, 4.4 * inch])
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f1f5f9")),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#111827")),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5f5")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.2 * inch))

    training = payload.get("training")
    if training:
        story.append(Paragraph("Entrenamiento", subtitle_style))
        story.append(Paragraph(str(training), normal_style))

    diet = payload.get("diet")
    if diet:
        story.append(Paragraph("Enfoque dietario", subtitle_style))
        story.append(Paragraph(str(diet), normal_style))

    calories = payload.get("calories")
    macros = payload.get("macros")
    if calories or macros:
        story.append(Paragraph("Macros y calorias", subtitle_style))
        if calories:
            story.append(Paragraph(f"Calorias: {calories}", normal_style))
        if macros:
            story.append(Paragraph(f"Macros: {macros}", normal_style))

    meals = payload.get("meals") or []
    if meals:
        story.append(Paragraph("Ejemplos de comidas", subtitle_style))
        for meal in meals:
            story.append(Paragraph(f"• {meal}", normal_style))

    sources = payload.get("sources") or []
    if sources:
        story.append(Paragraph("Fuentes", subtitle_style))
        for src in sources:
            story.append(Paragraph(str(src), normal_style))

    footer_style = ParagraphStyle(
        "HeroFooter",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#94a3b8"),
        alignment=TA_CENTER,
    )
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(f"Generado por Fitter • {datetime.now().strftime('%d/%m/%Y %H:%M')}", footer_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_routine_docx(routine_data: Dict[str, Any]) -> BytesIO:
    """
    Genera un documento Word de la rutina de entrenamiento.
    
    Args:
        routine_data: Diccionario con la estructura de routine_detail del chatbot
        
    Returns:
        BytesIO con el contenido del documento Word
    """
    buffer = BytesIO()
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Título principal
    header = routine_data.get("header", "Rutina de Entrenamiento")
    title = doc.add_heading(header, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(30, 64, 175)
    
    # Información del resumen
    doc.add_heading('Resumen', level=1)
    summary = routine_data.get("summary", {})
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    summary_items = [
        ("Tiempo estimado:", f"{summary.get('tiempo_min', '-')} minutos"),
        ("Ejercicios:", str(summary.get('ejercicios', '-'))),
        ("Equipamiento:", summary.get('equipamiento', '-')),
        ("Objetivo:", summary.get('objetivo', '-')),
        ("Nivel:", summary.get('nivel', '-')),
        ("Músculo:", summary.get('musculo', '-')),
    ]
    
    for idx, (label, value) in enumerate(summary_items):
        row = table.rows[idx]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    
    # Precauciones de salud
    health_notes = summary.get('health_notes')
    if health_notes and isinstance(health_notes, list) and len(health_notes) > 0:
        doc.add_heading('⚠️ Precauciones', level=2)
        for note in health_notes:
            p = doc.add_paragraph(note, style='List Bullet')
            p.runs[0].font.color.rgb = RGBColor(185, 28, 28)
    
    # Alergias
    allergies = summary.get('allergies')
    if allergies and isinstance(allergies, list) and len(allergies) > 0:
        p = doc.add_paragraph()
        p.add_run(f"Alergias registradas: {', '.join(allergies)}")
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(10)
    
    # Fallback notice
    fallback_notice = routine_data.get('fallback_notice')
    if fallback_notice:
        p = doc.add_paragraph()
        p.add_run(f"ℹ️ {fallback_notice}")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(146, 64, 14)
    
    # Ejercicios
    exercises = routine_data.get('exercises', [])
    if exercises:
        doc.add_heading('Detalle de Ejercicios', level=1)
        
        for exercise in exercises:
            ex_name = exercise.get('nombre', 'Ejercicio')
            series = exercise.get('series', '-')
            reps = exercise.get('repeticiones', '-')
            rpe = exercise.get('rpe', '-')
            rir = exercise.get('rir', '-')
            video = exercise.get('video', '')
            orden = exercise.get('orden', '')
            
            # Nombre del ejercicio
            p = doc.add_paragraph()
            run = p.add_run(f"{orden}. {ex_name}")
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(31, 41, 55)
            
            # Detalles
            details = f"Series: {series} • Repeticiones: {reps} • {rpe} • {rir}"
            p = doc.add_paragraph(details)
            
            if video:
                p = doc.add_paragraph()
                run = p.add_run(f"🎥 Video: {video}")
                run.font.color.rgb = RGBColor(37, 99, 235)
                run.font.size = Pt(9)
    
    # Progresión
    progresion = summary.get('progresion')
    if progresion:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run(f"📈 Progresión: {progresion}")
        run.font.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Generado por Fitter • {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(156, 163, 175)
    
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_diet_pdf(diet_data: Dict[str, Any]) -> BytesIO:
    """
    Genera un PDF de la dieta.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch
    )
    story = []
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, alignment=TA_CENTER)
    normal = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=10)

    header = diet_data.get('header', 'Dieta Fitter')
    story.append(Paragraph(header, title_style))
    story.append(Spacer(1, 0.2 * inch))

    summary = diet_data.get('summary', {})
    summary_items = []
    if 'target_kcal' in summary:
        summary_items.append(['Kcal objetivo:', str(summary.get('target_kcal'))])
    if 'proteinas_g' in summary:
        summary_items.append(['Proteínas (g):', str(summary.get('proteinas_g'))])
    if 'carbs_g' in summary:
        summary_items.append(['Carbohidratos (g):', str(summary.get('carbs_g'))])
    if 'fats_g' in summary:
        summary_items.append(['Grasas (g):', str(summary.get('fats_g'))])

    if summary_items:
        table = Table(summary_items, colWidths=[2.5 * inch, 3.5 * inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#111827')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#d1d5db')),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.2 * inch))

    meals = diet_data.get('meals', [])
    if meals:
        for meal in meals:
            meal_title = Paragraph(meal.get('name', 'Comida'), styles['Heading2'])
            story.append(meal_title)
            story.append(Spacer(1, 0.05 * inch))
            for item in meal.get('items', []):
                # Handle both string format and object format
                if isinstance(item, str):
                    line = f"- {item}"
                elif isinstance(item, dict):
                    name = item.get('name') or item.get('descripcion') or 'Alimento'
                    qty = item.get('qty', '')
                    kcal = item.get('kcal')
                    line = f"- {name}"
                    if qty:
                        line += f" {qty}"
                    if kcal:
                        line += f" • {kcal} kcal"
                else:
                    line = f"- {str(item)}"
                story.append(Paragraph(line, normal))
            story.append(Spacer(1, 0.15 * inch))

    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER)
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(f"Generado por Fitter • {datetime.now().strftime('%d/%m/%Y %H:%M')}", footer_style))

    doc.build(story)
    buffer.seek(0)
    return buffer


def generate_diet_docx(diet_data: Dict[str, Any]) -> BytesIO:
    buffer = BytesIO()
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    doc.add_heading(diet_data.get('header', 'Dieta Fitter'), level=0)

    summary = diet_data.get('summary', {})
    if summary:
        doc.add_heading('Resumen', level=1)
        for k, v in [('Kcal objetivo', summary.get('target_kcal')), ('Proteínas (g)', summary.get('proteinas_g')), ('Carbs (g)', summary.get('carbs_g')), ('Grasas (g)', summary.get('fats_g'))]:
            if v is not None:
                p = doc.add_paragraph()
                p.add_run(f"{k}: ").bold = True
                p.add_run(str(v))

    meals = diet_data.get('meals', [])
    if meals:
        for meal in meals:
            doc.add_heading(meal.get('name', 'Comida'), level=2)
            for item in meal.get('items', []):
                # Handle both string format and object format
                if isinstance(item, str):
                    text = f"- {item}"
                elif isinstance(item, dict):
                    name = item.get('name') or item.get('descripcion') or 'Alimento'
                    qty = item.get('qty', '')
                    kcal = item.get('kcal')
                    text = f"- {name}"
                    if qty:
                        text += f" {qty}"
                    if kcal:
                        text += f" • {kcal} kcal"
                else:
                    text = f"- {str(item)}"
                doc.add_paragraph(text)

    doc.add_paragraph()
    doc.add_paragraph(f"Generado por Fitter • {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.save(buffer)
    buffer.seek(0)
    return buffer
